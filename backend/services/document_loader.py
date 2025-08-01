"""
Unified Document Loader Service
Handles robust loading of multiple document formats with error recovery
Supports: CSV, JSON, Excel, PDF, DOCX, Markdown, Text, Images
"""

import os
import json
import pandas as pd
import numpy as np
from pathlib import Path
from typing import Dict, Any, List, Optional, Union, Tuple
import mimetypes
import chardet
import asyncio
from datetime import datetime
import hashlib
import io

# Import format-specific libraries with fallbacks
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    print("[DocumentLoader] PyPDF2 not available")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    print("[DocumentLoader] openpyxl not available")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[DocumentLoader] python-docx not available")

from PIL import Image

class DocumentLoadError(Exception):
    """Custom exception for document loading errors"""
    pass

class UnifiedDocumentLoader:
    """
    Unified document loader with robust error handling and format detection
    """
    
    # Supported formats and their handlers
    FORMAT_HANDLERS = {
        'csv': ['load_csv', 'load_csv_fallback'],
        'json': ['load_json', 'load_json_fallback'],
        'xlsx': ['load_excel', 'load_excel_fallback'],
        'xls': ['load_excel', 'load_excel_fallback'],
        'pdf': ['load_pdf_native', 'load_pdf_fallback'],
        'docx': ['load_docx', 'load_docx_fallback'],
        'doc': ['load_docx', 'load_docx_fallback'],
        'txt': ['load_text', 'load_text_fallback'],
        'md': ['load_text', 'load_text_fallback'],
        'rtf': ['load_text', 'load_text_fallback'],
        'png': ['load_image', 'load_image_fallback'],
        'jpg': ['load_image', 'load_image_fallback'],
        'jpeg': ['load_image', 'load_image_fallback'],
        'gif': ['load_image', 'load_image_fallback'],
        'webp': ['load_image', 'load_image_fallback'],
    }
    
    # Common encodings to try
    ENCODINGS = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16', 'utf-32']
    
    def __init__(self, claude_service=None):
        self.claude_service = claude_service
        self.cache = {}  # Simple in-memory cache
        
    async def load_document(
        self, 
        file_path: Union[str, Path], 
        file_type: Optional[str] = None,
        options: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """
        Load a document with automatic format detection and error recovery
        
        Args:
            file_path: Path to the document
            file_type: Optional MIME type or extension
            options: Loading options (encoding, delimiter, etc.)
            
        Returns:
            Dictionary containing:
            - content: Raw text content
            - structured_data: DataFrame for structured formats
            - metadata: Document metadata
            - format: Detected format
            - processing_method: Method used to load
            - warnings: Any warnings during loading
        """
        file_path = Path(file_path)
        options = options or {}
        
        # Check cache
        cache_key = self._get_cache_key(file_path)
        if cache_key in self.cache and not options.get('force_reload'):
            print(f"[DocumentLoader] Using cached version for {file_path.name}")
            return self.cache[cache_key]
        
        # Detect file format
        detected_format = self._detect_format(file_path, file_type)
        print(f"[DocumentLoader] Loading {file_path.name} as {detected_format}")
        
        # Get handlers for this format
        handlers = self.FORMAT_HANDLERS.get(detected_format, ['load_generic'])
        
        # Try each handler in order
        last_error = None
        warnings = []
        
        for handler_name in handlers:
            try:
                handler = getattr(self, handler_name, None)
                if handler:
                    result = await handler(file_path, options)
                    result['format'] = detected_format
                    result['processing_method'] = handler_name
                    result['warnings'] = warnings
                    
                    # Cache result
                    self.cache[cache_key] = result
                    
                    return result
            except Exception as e:
                last_error = e
                warnings.append(f"{handler_name} failed: {str(e)}")
                print(f"[DocumentLoader] {handler_name} failed for {file_path.name}: {e}")
                continue
        
        # All handlers failed
        raise DocumentLoadError(
            f"Failed to load {file_path.name} as {detected_format}. "
            f"Last error: {last_error}"
        )
    
    def _detect_format(self, file_path: Path, file_type: Optional[str] = None) -> str:
        """Detect file format from extension or MIME type"""
        if file_type:
            # Map MIME types to our format keys
            mime_map = {
                'text/csv': 'csv',
                'application/csv': 'csv',
                'application/json': 'json',
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
                'application/vnd.ms-excel': 'xls',
                'application/pdf': 'pdf',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                'application/msword': 'doc',
                'text/plain': 'txt',
                'text/markdown': 'md',
                'application/rtf': 'rtf',
                'image/png': 'png',
                'image/jpeg': 'jpg',
                'image/jpg': 'jpg',
                'image/gif': 'gif',
                'image/webp': 'webp',
            }
            format_key = mime_map.get(file_type)
            if format_key:
                return format_key
        
        # Fallback to extension
        extension = file_path.suffix.lower().lstrip('.')
        return extension or 'unknown'
    
    def _get_cache_key(self, file_path: Path) -> str:
        """Generate cache key based on file path and modification time"""
        stat = file_path.stat()
        key_string = f"{file_path}:{stat.st_mtime}:{stat.st_size}"
        return hashlib.md5(key_string.encode()).hexdigest()
    
    def _detect_encoding(self, file_path: Path, sample_size: int = 10000) -> str:
        """Detect file encoding using chardet"""
        try:
            with open(file_path, 'rb') as f:
                sample = f.read(sample_size)
                result = chardet.detect(sample)
                confidence = result.get('confidence', 0)
                encoding = result.get('encoding', 'utf-8')
                
                if confidence > 0.8:
                    return encoding
                else:
                    print(f"[DocumentLoader] Low confidence encoding detection: {encoding} ({confidence})")
                    return 'utf-8'
        except Exception:
            return 'utf-8'
    
    async def load_csv(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Load CSV file with robust error handling"""
        encoding = options.get('encoding') or self._detect_encoding(file_path)
        delimiter = options.get('delimiter')
        
        # Try to detect delimiter if not provided
        if not delimiter:
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                sample = f.read(1000)
                # Common delimiters
                for delim in [',', '\t', ';', '|']:
                    if delim in sample:
                        delimiter = delim
                        break
                else:
                    delimiter = ','
        
        # Load with pandas
        df = pd.read_csv(
            file_path,
            encoding=encoding,
            delimiter=delimiter,
            on_bad_lines='warn',
            encoding_errors='replace',
            low_memory=False
        )
        
        # Generate content summary
        content = f"CSV with {len(df)} rows and {len(df.columns)} columns\n"
        content += f"Columns: {', '.join(df.columns)}\n"
        content += df.head(10).to_string()
        
        return {
            'content': content,
            'structured_data': df,
            'metadata': {
                'rows': len(df),
                'columns': list(df.columns),
                'dtypes': {col: str(dtype) for col, dtype in df.dtypes.items()},
                'encoding': encoding,
                'delimiter': delimiter
            }
        }
    
    async def load_csv_fallback(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback CSV loader with manual parsing"""
        # Try different encodings
        content = None
        encoding_used = None
        
        for encoding in self.ENCODINGS:
            try:
                with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                    content = f.read()
                    encoding_used = encoding
                    break
            except Exception:
                continue
        
        if not content:
            raise DocumentLoadError("Unable to read file with any encoding")
        
        # Manual CSV parsing
        lines = content.strip().split('\n')
        if not lines:
            raise DocumentLoadError("Empty CSV file")
        
        # Simple CSV parsing (handles basic cases)
        rows = []
        for line in lines:
            # Handle quoted values
            if '"' in line:
                # Basic quoted value handling
                parts = []
                current = ""
                in_quotes = False
                for char in line:
                    if char == '"':
                        in_quotes = not in_quotes
                    elif char == ',' and not in_quotes:
                        parts.append(current.strip())
                        current = ""
                    else:
                        current += char
                parts.append(current.strip())
                rows.append(parts)
            else:
                rows.append([cell.strip() for cell in line.split(',')])
        
        # Create DataFrame
        if len(rows) > 1:
            df = pd.DataFrame(rows[1:], columns=rows[0])
        else:
            df = pd.DataFrame(rows)
        
        return {
            'content': content[:5000],  # First 5000 chars
            'structured_data': df,
            'metadata': {
                'rows': len(df),
                'columns': list(df.columns) if len(rows) > 1 else [],
                'encoding': encoding_used,
                'fallback_method': 'manual_parsing'
            }
        }
    
    async def load_json(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Load JSON file with support for nested structures"""
        encoding = options.get('encoding', 'utf-8')
        
        with open(file_path, 'r', encoding=encoding) as f:
            data = json.load(f)
        
        # Handle different JSON structures
        if isinstance(data, list):
            # Array of objects
            if data and isinstance(data[0], dict):
                df = pd.json_normalize(data, sep='_')
                content = f"JSON array with {len(data)} objects\n"
                content += f"Keys: {list(data[0].keys())}\n"
                content += json.dumps(data[:3], indent=2)
            else:
                # Array of primitives
                df = pd.DataFrame(data, columns=['value'])
                content = f"JSON array with {len(data)} values\n"
                content += str(data[:10])
        elif isinstance(data, dict):
            # Single object or nested structure
            try:
                # Try to normalize nested structure
                df = pd.json_normalize(data, sep='_')
                if len(df) == 1:
                    # Single object - might want to transpose
                    df = df.T.reset_index()
                    df.columns = ['key', 'value']
            except Exception:
                # Fallback to key-value pairs
                df = pd.DataFrame(list(data.items()), columns=['key', 'value'])
            
            content = f"JSON object with {len(data)} keys\n"
            content += json.dumps(data, indent=2)[:2000]
        else:
            # Primitive value
            df = pd.DataFrame([{'value': data}])
            content = f"JSON value: {data}"
        
        return {
            'content': content,
            'structured_data': df,
            'metadata': {
                'json_type': type(data).__name__,
                'keys': list(data.keys()) if isinstance(data, dict) else None,
                'length': len(data) if isinstance(data, (list, dict)) else 1
            }
        }
    
    async def load_json_fallback(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback JSON loader for corrupted files"""
        content = None
        
        # Try different encodings
        for encoding in self.ENCODINGS:
            try:
                with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                    content = f.read()
                    
                # Try to fix common JSON errors
                # Remove trailing commas
                import re
                content = re.sub(r',\s*}', '}', content)
                content = re.sub(r',\s*]', ']', content)
                
                # Try to parse
                data = json.loads(content)
                
                # Success - use normal JSON processing
                return await self.load_json(file_path, {'encoding': encoding})
                
            except json.JSONDecodeError as e:
                # Try to extract partial JSON
                try:
                    # Find the first complete JSON object/array
                    for i in range(len(content)):
                        try:
                            data = json.loads(content[:i+1])
                            df = pd.DataFrame([{'partial_data': str(data), 'error': str(e)}])
                            return {
                                'content': content[:1000],
                                'structured_data': df,
                                'metadata': {
                                    'error': 'Partial JSON extraction',
                                    'encoding': encoding
                                }
                            }
                        except:
                            continue
                except:
                    pass
            except Exception:
                continue
        
        # Complete failure - return raw content
        return {
            'content': content[:5000] if content else "",
            'structured_data': pd.DataFrame([{'raw_content': content[:1000] if content else ""}]),
            'metadata': {
                'error': 'Failed to parse JSON',
                'fallback': True
            }
        }
    
    async def load_excel(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Load Excel file with multi-sheet support"""
        sheet_name = options.get('sheet_name', None)
        
        # Load all sheets or specific sheet
        if sheet_name is None:
            # Load all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets = {}
            for sheet in excel_file.sheet_names:
                sheets[sheet] = pd.read_excel(file_path, sheet_name=sheet)
            
            # Combine sheets or use first sheet
            if len(sheets) == 1:
                df = list(sheets.values())[0]
                content = f"Excel file with 1 sheet: {list(sheets.keys())[0]}\n"
            else:
                # For multiple sheets, combine with sheet indicator
                dfs = []
                for sheet_name, sheet_df in sheets.items():
                    sheet_df['_sheet'] = sheet_name
                    dfs.append(sheet_df)
                df = pd.concat(dfs, ignore_index=True)
                content = f"Excel file with {len(sheets)} sheets: {', '.join(sheets.keys())}\n"
            
            content += f"Total rows: {len(df)}, columns: {len(df.columns)}"
            
            metadata = {
                'sheets': list(sheets.keys()),
                'total_rows': len(df),
                'columns': list(df.columns)
            }
        else:
            # Load specific sheet
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            content = f"Excel sheet '{sheet_name}' with {len(df)} rows and {len(df.columns)} columns"
            metadata = {
                'sheet': sheet_name,
                'rows': len(df),
                'columns': list(df.columns)
            }
        
        return {
            'content': content,
            'structured_data': df,
            'metadata': metadata
        }
    
    async def load_excel_fallback(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback Excel loader using openpyxl directly"""
        if not OPENPYXL_AVAILABLE:
            raise DocumentLoadError("openpyxl not available for Excel fallback")
        
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        
        # Get first sheet or specified sheet
        sheet_name = options.get('sheet_name', wb.sheetnames[0])
        ws = wb[sheet_name]
        
        # Read data
        data = []
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):  # Skip empty rows
                data.append(list(row))
        
        # Create DataFrame
        if data:
            df = pd.DataFrame(data[1:], columns=data[0] if len(data) > 1 else None)
        else:
            df = pd.DataFrame()
        
        wb.close()
        
        return {
            'content': f"Excel file loaded with openpyxl fallback\nSheet: {sheet_name}\nRows: {len(df)}",
            'structured_data': df,
            'metadata': {
                'sheet': sheet_name,
                'all_sheets': wb.sheetnames,
                'fallback_method': 'openpyxl'
            }
        }
    
    async def load_pdf_native(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Load PDF using Claude's native PDF support"""
        if not self.claude_service or not self.claude_service.supports_pdf_native:
            raise DocumentLoadError("Claude service not available for native PDF processing")
        
        # For native PDF processing, we just need to mark it for Claude
        # The actual processing will happen when we send it to Claude
        
        # Get basic metadata using PyPDF2
        metadata = {}
        text_preview = ""
        
        if PYPDF2_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    metadata['page_count'] = len(reader.pages)
                    metadata['info'] = {
                        'title': reader.metadata.get('/Title', '') if reader.metadata else '',
                        'author': reader.metadata.get('/Author', '') if reader.metadata else '',
                    }
                    
                    # Get first page text as preview
                    if reader.pages:
                        text_preview = reader.pages[0].extract_text()[:500]
            except Exception as e:
                print(f"[DocumentLoader] Error reading PDF metadata: {e}")
        
        return {
            'content': text_preview or "PDF document (will be processed with Claude Vision)",
            'structured_data': None,  # Will be extracted by Claude
            'metadata': metadata,
            'requires_claude_processing': True,
            'file_path': str(file_path)  # Keep path for Claude processing
        }
    
    async def load_pdf_fallback(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback PDF loader using PyPDF2"""
        if not PYPDF2_AVAILABLE:
            raise DocumentLoadError("PyPDF2 not available for PDF fallback")
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            # Extract text from all pages
            text_content = ""
            for i, page in enumerate(reader.pages):
                text_content += f"\n--- Page {i+1} ---\n"
                text_content += page.extract_text()
            
            # Create simple DataFrame with page content
            pages_data = []
            for i, page in enumerate(reader.pages):
                pages_data.append({
                    'page': i + 1,
                    'content': page.extract_text()
                })
            
            df = pd.DataFrame(pages_data) if pages_data else pd.DataFrame()
            
            return {
                'content': text_content,
                'structured_data': df,
                'metadata': {
                    'page_count': len(reader.pages),
                    'extraction_method': 'PyPDF2',
                    'has_text': bool(text_content.strip())
                }
            }
    
    async def load_docx(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Load Word document"""
        if not DOCX_AVAILABLE:
            raise DocumentLoadError("python-docx not available")
        
        doc = docx.Document(str(file_path))
        
        # Extract text content
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract tables
        tables_data = []
        for i, table in enumerate(doc.tables):
            table_content = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_content.append(row_data)
            
            if table_content:
                # Create DataFrame for each table
                table_df = pd.DataFrame(table_content[1:], columns=table_content[0] if len(table_content) > 1 else None)
                tables_data.append({
                    'table_index': i,
                    'data': table_df
                })
        
        # Combine content
        content = '\n\n'.join(paragraphs)
        
        # Create structured data
        if tables_data:
            # If we have tables, include them in structured data
            df = pd.DataFrame({
                'content_type': ['paragraph'] * len(paragraphs) + ['table'] * len(tables_data),
                'content': paragraphs + [f"Table {t['table_index']}" for t in tables_data]
            })
        else:
            df = pd.DataFrame({'paragraph': paragraphs}) if paragraphs else pd.DataFrame()
        
        return {
            'content': content,
            'structured_data': df,
            'metadata': {
                'paragraphs': len(paragraphs),
                'tables': len(tables_data),
                'core_properties': {
                    'author': doc.core_properties.author or '',
                    'title': doc.core_properties.title or '',
                    'created': doc.core_properties.created.isoformat() if doc.core_properties.created else ''
                }
            },
            'tables': tables_data  # Store tables separately
        }
    
    async def load_docx_fallback(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback for DOCX - extract raw XML"""
        # DOCX files are ZIP archives containing XML
        import zipfile
        
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Extract main document content
                with zip_file.open('word/document.xml') as xml_file:
                    content = xml_file.read().decode('utf-8')
                
                # Basic text extraction from XML
                import re
                # Remove XML tags (basic approach)
                text = re.sub(r'<[^>]+>', ' ', content)
                text = re.sub(r'\s+', ' ', text).strip()
                
                return {
                    'content': text[:5000],  # First 5000 chars
                    'structured_data': pd.DataFrame([{'text': text[:1000]}]),
                    'metadata': {
                        'extraction_method': 'xml_fallback',
                        'warning': 'Formatting and structure lost'
                    }
                }
        except Exception as e:
            raise DocumentLoadError(f"Failed to extract DOCX content: {e}")
    
    async def load_text(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Load text-based files (TXT, MD, RTF)"""
        encoding = options.get('encoding') or self._detect_encoding(file_path)
        
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            content = f.read()
        
        # Detect file type for metadata
        suffix = file_path.suffix.lower()
        file_type = 'markdown' if suffix == '.md' else 'text'
        
        # Create simple DataFrame with lines or paragraphs
        lines = content.split('\n')
        paragraphs = content.split('\n\n')
        
        df = pd.DataFrame({
            'line_number': range(1, len(lines) + 1),
            'content': lines
        })
        
        return {
            'content': content,
            'structured_data': df,
            'metadata': {
                'encoding': encoding,
                'file_type': file_type,
                'lines': len(lines),
                'paragraphs': len(paragraphs),
                'characters': len(content)
            }
        }
    
    async def load_text_fallback(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback text loader with binary mode"""
        # Read as binary and try to decode
        with open(file_path, 'rb') as f:
            raw_content = f.read()
        
        # Try multiple decodings
        content = None
        encoding_used = None
        
        for encoding in self.ENCODINGS + ['ascii', 'utf-8-sig']:
            try:
                content = raw_content.decode(encoding)
                encoding_used = encoding
                break
            except Exception:
                continue
        
        if content is None:
            # Last resort - decode with replacement
            content = raw_content.decode('utf-8', errors='replace')
            encoding_used = 'utf-8 (with replacements)'
        
        return {
            'content': content,
            'structured_data': pd.DataFrame([{'text': content[:1000]}]),
            'metadata': {
                'encoding': encoding_used,
                'size_bytes': len(raw_content),
                'fallback_method': 'binary_decode'
            }
        }
    
    async def load_image(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Load image files"""
        # Open image to get metadata
        img = Image.open(file_path)
        
        # Get image info
        metadata = {
            'format': img.format,
            'mode': img.mode,
            'size': img.size,
            'width': img.width,
            'height': img.height
        }
        
        # For images, we'll need Claude's vision API for content extraction
        return {
            'content': f"Image file: {file_path.name} ({img.width}x{img.height} {img.format})",
            'structured_data': pd.DataFrame([metadata]),
            'metadata': metadata,
            'requires_claude_processing': True,
            'file_path': str(file_path)
        }
    
    async def load_image_fallback(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback image loader"""
        # Just return basic file info
        stat = file_path.stat()
        
        return {
            'content': f"Image file: {file_path.name}",
            'structured_data': pd.DataFrame([{
                'filename': file_path.name,
                'size_bytes': stat.st_size,
                'modified': datetime.fromtimestamp(stat.st_mtime).isoformat()
            }]),
            'metadata': {
                'size_bytes': stat.st_size,
                'fallback_method': 'file_stats_only'
            }
        }
    
    async def load_generic(self, file_path: Path, options: Dict[str, Any]) -> Dict[str, Any]:
        """Generic loader for unknown file types"""
        # Try to read as text
        try:
            return await self.load_text(file_path, options)
        except Exception:
            # Return file info only
            stat = file_path.stat()
            
            return {
                'content': f"Unknown file type: {file_path.name}",
                'structured_data': pd.DataFrame([{
                    'filename': file_path.name,
                    'size_bytes': stat.st_size,
                    'modified': datetime.fromtimestamp(stat.st_mtime).isoformat()
                }]),
                'metadata': {
                    'size_bytes': stat.st_size,
                    'mime_type': mimetypes.guess_type(str(file_path))[0],
                    'warning': 'Unknown file type'
                }
            }
    
    async def load_multiple_documents(
        self, 
        file_paths: List[Union[str, Path]], 
        options: Dict[str, Any] = None
    ) -> List[Dict[str, Any]]:
        """
        Load multiple documents in parallel
        
        Args:
            file_paths: List of file paths
            options: Loading options
            
        Returns:
            List of loaded documents
        """
        options = options or {}
        
        # Create tasks for parallel loading
        tasks = []
        for file_path in file_paths:
            task = self.load_document(file_path, options=options)
            tasks.append(task)
        
        # Execute in parallel
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Process results
        loaded_docs = []
        errors = []
        
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                errors.append({
                    'file': str(file_paths[i]),
                    'error': str(result)
                })
                print(f"[DocumentLoader] Error loading {file_paths[i]}: {result}")
            else:
                loaded_docs.append(result)
        
        if errors:
            print(f"[DocumentLoader] Failed to load {len(errors)} documents")
        
        return loaded_docs

# Global instance
document_loader = UnifiedDocumentLoader()